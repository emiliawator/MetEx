# docx extension

import docx
from datetime import datetime

def read(filepath):
    try:
        doc = docx.Document(filepath)
    except:
        return "Supplied filepath is invalid"
    try:
        prop = doc.core_properties
    except: 
        return "An error occured while reading metadata"
                
    metadata = [] # list of tuples

    # str if not date
    metadata.append(tuple(("category", prop.category)))
    metadata.append(tuple(("content_status", prop.content_status)))
    metadata.append(tuple(("created", prop.created))) #date
    metadata.append(tuple(("author", prop.author)))
    metadata.append(tuple(("comments", prop.comments)))
    metadata.append(tuple(("identifier", prop.identifier)))
    metadata.append(tuple(("keywords", prop.keywords)))
    metadata.append(tuple(("language", prop.language)))
    metadata.append(tuple(("last_modified_by", prop.last_modified_by)))
    metadata.append(tuple(("last_printed", prop.last_printed))) #date
    metadata.append(tuple(("modified", prop.modified))) #date
    metadata.append(tuple(("revision", prop.revision))) # int
    metadata.append(tuple(("subject", prop.subject)))
    metadata.append(tuple(("title", prop.title)))
    metadata.append(tuple(("version", prop.version)))
    
    return metadata
 

def save(metadata, filepath, newfilepath):
    errors = []
    try:
        doc = docx.Document(filepath)
    except:
        errors.append("Supplied filepath is invalid")
        return errors
    try:
        prop = doc.core_properties
    except:
        errors.append("An error occured while reading metadata")
        return errors

    notuplelist = [list(i) for i in metadata]

    for item in notuplelist:
        if item[0] == "created" or item[0] == "last_printed" or item[0] == "modified": # datetime type
            if item[1] != "None":
                    try:
                        item[1] = datetime.strptime(item[1], "%Y-%m-%d %H:%M:%S")
                    except:
                        errors.append("Correct date format must be supplied in {} field: \n %Y-%m-%d %H:%M:%S".format(item[0]))
                        return errors
            else: # if None by default
                continue
            try:
                setattr(prop, item[0], item[1])
            except:
                errors.append("An error occured while saving {} field".format(item[0]))
                return errors
            
        if item[0] == "revision": # positive int
            try:
                setattr(prop, item[0], (int)(item[1]))
            except:
                errors.append("Revision field must be a positive integer value")
                return errors
        else:
            try:
                setattr(prop, item[0], item[1])
            except:
                errors.append("An error occured while saving {} field".format(item[0]))
                return errors
    try:
        doc.save(newfilepath)
    except:
        errors.append("Supplied filepath is invalid")
        return errors
    
def erase(filepath):
    errors = []
    try:
        doc = docx.Document(filepath)
    except:
        errors.append("Supplied filepath is invalid")
        return errors
    try:
        prop = doc.core_properties
    except:
        errors.append("An error occured while reading metadata")
        return errors
    metadata = [] # list of tuples

     
    # str if not date
    metadata.append(tuple(("category", "")))
    metadata.append(tuple(("content_status", "")))
    metadata.append(tuple(("author", "")))
    metadata.append(tuple(("comments", "")))
    metadata.append(tuple(("identifier", "")))
    metadata.append(tuple(("keywords", "")))
    metadata.append(tuple(("language", "")))
    metadata.append(tuple(("last_modified_by", "")))
    metadata.append(tuple(("revision", 1))) # positive int
    metadata.append(tuple(("subject", "")))
    metadata.append(tuple(("title", "")))
    metadata.append(tuple(("version", "")))
    
    # datetime handling
    default_date = datetime.min
    if prop.created is not None:
        metadata.append(tuple(("created", default_date)))
    if prop.last_printed is not None:
        metadata.append(tuple(("last_printed", default_date)))
    if prop.modified is not None:
        metadata.append(tuple(("modified", default_date)))
    
    notuplelist = [list(i) for i in metadata]

    for item in notuplelist:
        setattr(prop, item[0], item[1])

    try:
        doc.save(filepath)
    except:
        errors.append("Supplied filepath is invalid")
        return errors